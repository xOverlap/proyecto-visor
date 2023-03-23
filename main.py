from io import BytesIO
from fastapi import FastAPI, File, UploadFile
from fastapi.responses import StreamingResponse
from fastapi.middleware.cors import CORSMiddleware
import os, zipfile

app = FastAPI()

origins = [
    "http://localhost",
    "http://localhost:5173",
]

app.add_middleware(
    CORSMiddleware,
    allow_origins=origins,
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Inicio conversor de PDF

@app.post("/api/v1/convert-pdf-to-csv")
async def convert_pdf_csv(file: UploadFile = File(...)):
    import tabula
    # Función para eliminar los archivos y carpetas creadas
    def deleteFiles():
        # Si la carpeta doc_files existe, elimina todos los archivos PDF y luego elimina la carpeta doc_files
        if os.path.exists("pdf_to_csv_files"):
            for pdf_file_ in os.listdir("pdf_to_csv_files"):
                os.remove(os.path.join("pdf_to_csv_files", pdf_file_))
            os.rmdir("pdf_to_csv_files")
        # Si la carpeta csv_files existe, elimina todos los archivos CSV y luego elimina la carpeta csv_files
        if os.path.exists("pdf_csv_files"):
            for pdf_file_ in os.listdir("pdf_csv_files"):
                a = open(os.path.join("pdf_csv_files", pdf_file_), "rb")
                a = a.close()

                os.remove(os.path.join("pdf_csv_files", pdf_file_))
            os.rmdir("pdf_csv_files")
    # Elimina los archivos y carpetas creadas en la última ejecución
    deleteFiles()
    # Crea una carpeta para almacenar el PDF recibido
    os.mkdir("pdf_to_csv_files")
    # Guardar el PDF recibido en el directorio doc_files
    file_name = os.path.join("pdf_to_csv_files", file.filename)
    # Abre el archivo PDF en modo binario y lo guarda en el directorio doc_files
    with open(file_name, "wb") as buffer:
        # Lee el archivo PDF en trozos de 1024 bytes
        while chunk := await file.read(1024):
            # Escribe el trozo de bytes en el archivo PDF
            buffer.write(chunk)
    # Especifica la ruta de tu archivo PDF
    file = file_name

    # Extrae todas las tablas del PDF
    all_tables = tabula.read_pdf(file, pages="all", multiple_tables=True)

    # Crea una carpeta para almacenar los archivos CSV extraídos
    os.makedirs("pdf_csv_files", exist_ok=True)

    # Utiliza un bucle for para ajustar el formato de cada tabla extraída y guardarla como un archivo CSV
    for i, table in enumerate(all_tables):
        # Elimina las filas y columnas vacías
        table.dropna(how="all", inplace=True)
        table.dropna(axis=1, how="all", inplace=True)
        # Reemplaza los nombres de las filas y columnas no deseadas con valores en blanco
        table.rename(columns=lambda x: "" if "Unnamed" in str(x) else x, inplace=True)
        table.rename(index=lambda x: "" if "Unnamed" in str(x) else x, inplace=True)
        # Si hay más de un archivo CSV, crea un archivo .zip y agrega todos los archivos CSV a él
        if len(all_tables) > 1:
            csv_file = os.path.join("pdf_csv_files", f"table_{i}.csv")
        else:
            # Si solo hay un archivo CSV, renombra el archivo con el nombre del PDF
            csv_file = os.path.join("pdf_csv_files", os.path.splitext(os.path.basename(file))[0] + ".csv")
        table.to_csv(csv_file, index=False)
    # Si hay más de un archivo CSV, crea un archivo .zip y agrega todos los archivos CSV a él
    zip_file_name = "csv_files.zip"
    with zipfile.ZipFile(zip_file_name, "w") as zip_file:
        # Agrega todos los archivos CSV a la carpeta csv_files
        for pdf_file_name in os.listdir("pdf_csv_files"):
            csv_file = os.path.join("pdf_csv_files", pdf_file_name)
            zip_file.write(csv_file, pdf_file_name)

    # Elimina la carpeta csv_files y sus contenidos
    for pdf_file_name in os.listdir("pdf_csv_files"):
        csv_file = os.path.join("pdf_csv_files", pdf_file_name)
        os.remove(csv_file)

    # Mueve el archivo zip a la carpeta csv_files
    os.rename("csv_files.zip", "pdf_csv_files/csv_files.zip")
    # Si el archivo zip existe, devuelve el archivo zip como respuesta y elimina los archivos y carpetas creadas
    if os.path.exists("pdf_csv_files/csv_files.zip"):
        data = open("pdf_csv_files/csv_files.zip", "rb")
        file_contents = data.read()
        data.close()
        response = StreamingResponse(BytesIO(file_contents), media_type="application/zip", headers={"Content-Disposition": "attachment; filename=csv_files.zip"})
        deleteFiles()
        return response

@app.post("/api/v1/convert-pdf-to-xls")
async def convert_pdf_xls(file: UploadFile = File(...)):
    import tabula
    import pandas as pd
    # Función para eliminar los archivos y carpetas creadas
    def deleteFiles():
        # Si la carpeta doc_files existe, elimina todos los archivos PDF y luego elimina la carpeta doc_files
        if os.path.exists("pdf_to_xls_files"):
            for pdf_file_ in os.listdir("pdf_to_xls_files"):
                os.remove(os.path.join("pdf_to_xls_files", pdf_file_))
            os.rmdir("pdf_to_xls_files")
        # Si la carpeta xls_files existe, elimina todos los archivos XLS y luego elimina la carpeta xls_files
        if os.path.exists("pdf_xls_files"):
            for xls_file_ in os.listdir("pdf_xls_files"):
                a = open(os.path.join("pdf_xls_files", xls_file_), "rb")
                a = a.close()

                os.remove(os.path.join("pdf_xls_files", xls_file_))
            os.rmdir("pdf_xls_files")
    # Elimina los archivos y carpetas creadas en la última ejecución
    deleteFiles()
    # Crea una carpeta para almacenar el PDF recibido
    os.mkdir("pdf_to_xls_files")
    # Guardar el PDF recibido en el directorio doc_files
    file_name = os.path.join("pdf_to_xls_files", file.filename)
    # Abre el archivo PDF en modo binario y lo guarda en el directorio doc_files
    with open(file_name, "wb") as buffer:
        # Lee el archivo PDF en trozos de 1024 bytes
        while chunk := await file.read(1024):
            # Escribe el trozo de bytes en el archivo PDF
            buffer.write(chunk)
    # Especifica la ruta de tu archivo PDF
    file = file_name

    # Extrae todas las tablas del PDF
    all_tables = tabula.read_pdf(file, pages="all", multiple_tables=True)

    # Crea una carpeta para almacenar los archivos XLS extraídos
    os.makedirs("pdf_xls_files", exist_ok=True)

    # Utiliza un bucle for para ajustar el formato de cada tabla extraída y guardarla como un archivo XLS
    for i, table in enumerate(all_tables):
        # Elimina las filas y columnas vacías
        table.dropna(how="all", inplace=True)
        table.dropna(axis=1, how="all", inplace=True)
        # Reemplaza los nombres de las filas y columnas no deseadas con valores en blanco
        table.rename(columns=lambda x: "" if "Unnamed" in str(x) else x, inplace=True)
        table.rename(index=lambda x: "" if "Unnamed" in str(x) else x, inplace=True)
        # Si hay más de un archivo XLS, crea un archivo .zip y agrega todos los archivos XLS a él
        if len(all_tables) > 1:
            xls_file = os.path.join("pdf_xls_files", f"table_{i}.xls")
        else:
            # Si solo hay un archivo XLS, renombra el archivo con el nombre del PDF
            xls_file = os.path.join("pdf_xls_files", os.path.splitext(os.path.basename(file))[0] + ".xls")
        table.to_excel(xls_file, index=False)
    # Si hay más de un archivo XLS, crea un archivo .zip y agrega todos los archivos XLS a él
    zip_file_name = "xls_files.zip"
    with zipfile.ZipFile(zip_file_name, "w") as zip_file:
        # Agrega todos los archivos XLS a la carpeta xls_files
        for xls_file_name in os.listdir("pdf_xls_files"):
            xls_file = os.path.join("pdf_xls_files", xls_file_name)
            zip_file.write(xls_file, xls_file_name)

    # Elimina la carpeta xls_files y sus contenidos
    for xls_file_name in os.listdir("pdf_xls_files"):
        xls_file = os.path.join("pdf_xls_files", xls_file_name)
        os.remove(xls_file)

    # Mueve el archivo zip a la carpeta xls_files
    os.rename("xls_files.zip", "pdf_xls_files/xls_files.zip")
    # Si el archivo zip existe, devuelve el archivo zip como respuesta y elimina los archivos y carpetas creadas
    if os.path.exists("pdf_xls_files/xls_files.zip"):
        data = open("pdf_xls_files/xls_files.zip", "rb")
        file_contents = data.read()
        data.close()
        response = StreamingResponse(BytesIO(file_contents), media_type="application/zip", headers={"Content-Disposition": "attachment; filename=xls_files.zip"})
        deleteFiles()
        return response
    
@app.post("/api/v1/convert-pdf-to-xlsx")
async def convert_pdf_xlsx(file: UploadFile = File(...)):
    import tabula
    import pandas as pd
    # Función para eliminar los archivos y carpetas creadas
    def deleteFiles():
        # Si la carpeta doc_files existe, elimina todos los archivos PDF y luego elimina la carpeta doc_files
        if os.path.exists("pdf_to_xls_files"):
            for pdf_file_ in os.listdir("pdf_to_xls_files"):
                os.remove(os.path.join("pdf_to_xls_files", pdf_file_))
            os.rmdir("pdf_to_xls_files")
        # Si la carpeta xls_files existe, elimina todos los archivos XLSX y luego elimina la carpeta xls_files
        if os.path.exists("pdf_xls_files"):
            for xls_file_ in os.listdir("pdf_xls_files"):
                a = open(os.path.join("pdf_xls_files", xls_file_), "rb")
                a = a.close()

                os.remove(os.path.join("pdf_xls_files", xls_file_))
            os.rmdir("pdf_xls_files")
    # Elimina los archivos y carpetas creadas en la última ejecución
    deleteFiles()
    # Crea una carpeta para almacenar el PDF recibido
    os.mkdir("pdf_to_xls_files")
    # Guardar el PDF recibido en el directorio doc_files
    file_name = os.path.join("pdf_to_xls_files", file.filename)
    # Abre el archivo PDF en modo binario y lo guarda en el directorio doc_files
    with open(file_name, "wb") as buffer:
        # Lee el archivo PDF en trozos de 1024 bytes
        while chunk := await file.read(1024):
            # Escribe el trozo de bytes en el archivo PDF
            buffer.write(chunk)
    # Especifica la ruta de tu archivo PDF
    file = file_name

    # Extrae todas las tablas del PDF
    all_tables = tabula.read_pdf(file, pages="all", multiple_tables=True)

    # Crea una carpeta para almacenar los archivos XLSX extraídos
    os.makedirs("pdf_xls_files", exist_ok=True)

    # Utiliza un bucle for para ajustar el formato de cada tabla extraída y guardarla como un archivo XLSX
    for i, table in enumerate(all_tables):
        # Elimina las filas y columnas vacías
        table.dropna(how="all", inplace=True)
        table.dropna(axis=1, how="all", inplace=True)
        # Reemplaza los nombres de las filas y columnas no deseadas con valores en blanco
        table.rename(columns=lambda x: "" if "Unnamed" in str(x) else x, inplace=True)
        table.rename(index=lambda x: "" if "Unnamed" in str(x) else x, inplace=True)
        # Si hay más de un archivo XLSX, crea un archivo .zip y agrega todos los archivos XLSX a él
        if len(all_tables) > 1:
            xlsx_file = os.path.join("pdf_xls_files", f"table_{i}.xlsx")
        else:
            # Si solo hay un archivo XLSX, renombra el archivo con el nombre del PDF
            xlsx_file = os.path.join("pdf_xls_files", os.path.splitext(os.path.basename(file))[0] + ".xlsx")
        table.to_excel(xlsx_file, index=False)
    # Si hay más de un archivo XLSX, crea un archivo .zip y agrega todos los archivos XLSX a él
    zip_file_name = "xls_files.zip"
    with zipfile.ZipFile(zip_file_name, "w") as zip_file:
        # Agrega todos los archivos XLSX a la carpeta xls_files
        for xls_file_name in os.listdir("pdf_xls_files"):
            xlsx_file = os.path.join("pdf_xls_files", xls_file_name)
            zip_file.write(xlsx_file, xls_file_name)

    # Elimina la carpeta xls_files y sus contenidos
    for xls_file_name in os.listdir("pdf_xls_files"):
        xlsx_file = os.path.join("pdf_xls_files", xls_file_name)
        os.remove(xlsx_file)

    # Mueve el archivo zip a la carpeta xls_files
    os.rename("xls_files.zip", "pdf_xls_files/xls_files.zip")
    # Si el archivo zip existe, devuelve el archivo zip como respuesta y elimina los archivos y carpetas creadas
    if os.path.exists("pdf_xls_files/xls_files.zip"):
        data = open("pdf_xls_files/xls_files.zip", "rb")
        file_contents = data.read()
        data.close()
        response = StreamingResponse(BytesIO(file_contents), media_type="application/zip", headers={"Content-Disposition": "attachment; filename=xls_files.zip"})
        deleteFiles()
        return response

@app.post("/api/v1/convert-pdf-to-docx")
async def convert_pdf_docx(file: UploadFile = File(...)):
    from pdf2docx import Converter
    def deleteFiles():
        # Si la carpeta doc_files existe, elimina todos los archivos PDF y luego elimina la carpeta doc_files
        if os.path.exists("pdf_to_docx_files"):
            for pdf_file_ in os.listdir("pdf_to_docx_files"):
                os.remove(os.path.join("pdf_to_docx_files", pdf_file_))
            os.rmdir("pdf_to_docx_files")
        # Si la carpeta csv_files existe, elimina todos los archivos CSV y luego elimina la carpeta csv_files
        if os.path.exists("pdf_docx_files"):
            for pdf_file_ in os.listdir("pdf_docx_files"):
                a = open(os.path.join("pdf_docx_files", pdf_file_), "rb")
                a = a.close()

                os.remove(os.path.join("pdf_docx_files", pdf_file_))
            os.rmdir("pdf_docx_files")
    deleteFiles()
    # Crea una carpeta para almacenar el PDF recibido
    os.mkdir("pdf_to_docx_files")
    # Guardar el PDF recibido en el directorio doc_files
    file_name = os.path.join("pdf_to_docx_files", file.filename)
    # Abre el archivo PDF en modo binario y lo guarda en el directorio doc_files
    with open(file_name, "wb") as buffer:
        # Lee el archivo PDF en trozos de 1024 bytes
        while chunk := await file.read(1024):
            # Escribe el trozo de bytes en el archivo PDF
            buffer.write(chunk)
    os.makedirs("pdf_docx_files", exist_ok=True)
    cv = Converter(file_name)
    cv.convert("pdf_docx_files/converted.docx", start=0, end=None)
    cv.close()
    zip_file_name = "docx_files.zip"
    with zipfile.ZipFile(zip_file_name, "w") as zip_file:
        # Agrega todos los archivos CSV a la carpeta csv_files
        for pdf_file_name in os.listdir("pdf_docx_files"):
            csv_file = os.path.join("pdf_docx_files", pdf_file_name)
            zip_file.write(csv_file, pdf_file_name)
    # Mueve el archivo zip a la carpeta csv_files
    os.rename("docx_files.zip", "pdf_docx_files/docx_files.zip")
    data = open("pdf_docx_files/docx_files.zip", "rb")
    file_contents = data.read()
    data.close()
    response = StreamingResponse(BytesIO(file_contents), media_type="application/docx", headers={"Content-Disposition": "attachment; filename=csv_files.zip"})
    deleteFiles()
    return response

@app.post("/api/v1/convert-pdf-to-doc")
async def convert_pdf_doc(file: UploadFile = File(...)):
    from pdf2docx import Converter
    def deleteFiles():
        # Si la carpeta doc_files existe, elimina todos los archivos PDF y luego elimina la carpeta doc_files
        if os.path.exists("pdf_to_doc_files"):
            for pdf_file_ in os.listdir("pdf_to_doc_files"):
                os.remove(os.path.join("pdf_to_doc_files", pdf_file_))
            os.rmdir("pdf_to_doc_files")
        # Si la carpeta csv_files existe, elimina todos los archivos CSV y luego elimina la carpeta csv_files
        if os.path.exists("pdf_doc_files"):
            for pdf_file_ in os.listdir("pdf_doc_files"):
                a = open(os.path.join("pdf_doc_files", pdf_file_), "rb")
                a = a.close()

                os.remove(os.path.join("pdf_doc_files", pdf_file_))
            os.rmdir("pdf_doc_files")
    deleteFiles()
    # Crea una carpeta para almacenar el PDF recibido
    os.mkdir("pdf_to_doc_files")
    # Guardar el PDF recibido en el directorio doc_files
    file_name = os.path.join("pdf_to_doc_files", file.filename)
    # Abre el archivo PDF en modo binario y lo guarda en el directorio doc_files
    with open(file_name, "wb") as buffer:
        # Lee el archivo PDF en trozos de 1024 bytes
        while chunk := await file.read(1024):
            # Escribe el trozo de bytes en el archivo PDF
            buffer.write(chunk)
    os.makedirs("pdf_doc_files", exist_ok=True)
    cv = Converter(file_name)
    cv.convert("pdf_doc_files/converted.doc", start=0, end=None)
    cv.close()
    zip_file_name = "doc_files.zip"
    with zipfile.ZipFile(zip_file_name, "w") as zip_file:
        # Agrega todos los archivos DOC a la carpeta pdf_doc_files
        for pdf_file_name in os.listdir("pdf_doc_files"):
            csv_file = os.path.join("pdf_doc_files", pdf_file_name)
            zip_file.write(csv_file, pdf_file_name)
    # Mueve el archivo zip a la carpeta csv_files
    os.rename("doc_files.zip", "pdf_doc_files/doc_files.zip")
    data = open("pdf_doc_files/doc_files.zip", "rb")
    file_contents = data.read()
    data.close()
    response = StreamingResponse(BytesIO(file_contents), media_type="application/doc", headers={"Content-Disposition": "attachment; filename=csv_files.zip"})
    deleteFiles()
    return response


# Final conversor de PDF


# Inicio conversor de CSV


@app.post("/api/v1/convert-csv-to-pdf")
async def convert_csv_pdf(file: UploadFile = File(...)):
    # imports
    import csv
    from fpdf import FPDF
    def deleteFiles():
        # Si la carpeta pdf_files existe, elimina todos los archivos PDF y luego elimina la carpeta pdf_files
        if os.path.exists("csv_to_pdf_files"):
            for pdf_file_ in os.listdir("csv_to_pdf_files"):
                os.remove(os.path.join("csv_to_pdf_files", pdf_file_))
            os.rmdir("csv_to_pdf_files")
        # Si la carpeta csv_files existe, elimina todos los archivos CSV y luego elimina la carpeta csv_files
        if os.path.exists("csv_pdf_files"):
            for pdf_file_ in os.listdir("csv_pdf_files"):
                a = open(os.path.join("csv_pdf_files", pdf_file_), "rb")
                a = a.close()

                os.remove(os.path.join("csv_pdf_files", pdf_file_))
            os.rmdir("csv_pdf_files")
    deleteFiles()
    # Crea una carpeta para almacenar el PDF recibido
    os.mkdir("csv_to_pdf_files")
    # Guardar el PDF recibido en el directorio pdf_files
    file_name = os.path.join("csv_to_pdf_files", file.filename)
    # Abre el archivo PDF en modo binario y lo guarda en el directorio pdf_files
    with open(file_name, "wb") as buffer:
        # Lee el archivo PDF en trozos de 1024 bytes
        while chunk := await file.read(1024):
            # Escribe el trozo de bytes en el archivo PDF
            buffer.write(chunk)
    os.makedirs("csv_pdf_files", exist_ok=True)
    
    # Crea un objeto PDF
    pdf = FPDF()

    # Añade una página
    pdf.add_page()

    # Define el tamaño y la fuente del texto
    pdf.set_font("Arial", size=12)

    # Abre el archivo CSV y lee los datos
    with open(file_name, "r") as csv_file:
        csv_reader = csv.reader(csv_file)

        # Itera por cada fila y coloca los datos en el PDF
        for row in csv_reader:
            for item in row:
                pdf.cell(40, 10, str(item))
            pdf.ln()

    # Guarda el archivo PDF
    pdf.output(os.path.join("csv_pdf_files", "archivo.pdf"))
    
    zip_file_name = "pdf_files.zip"
    with zipfile.ZipFile(zip_file_name, "w") as zip_file:
        # Agrega todos los archivos DOC a la carpeta csv_pdf_files
        for pdf_file_name in os.listdir("csv_pdf_files"):
            csv_file = os.path.join("csv_pdf_files", pdf_file_name)
            zip_file.write(csv_file, pdf_file_name)
    # Mueve el archivo zip a la carpeta csv_files
    os.rename("pdf_files.zip", "csv_pdf_files/pdf_files.zip")
    data = open("csv_pdf_files/pdf_files.zip", "rb")
    file_contents = data.read()
    data.close()
    response = StreamingResponse(BytesIO(file_contents), media_type="application/pdf", headers={"Content-Disposition": "attachment; filename=csv_files.zip"})
    deleteFiles()
    return response

@app.post("/api/v1/convert-csv-to-xlsx")
async def convert_csv_xlsx(file: UploadFile = File(...)):
    import csv
    from openpyxl import Workbook

    def deleteFiles():
        # Si la carpeta doc_files existe, elimina todos los archivos PDF y luego elimina la carpeta doc_files
        if os.path.exists("csv_to_xlsx_files"):
            for csv_file_ in os.listdir("csv_to_xlsx_files"):
                os.remove(os.path.join("csv_to_xlsx_files", csv_file_))
            os.rmdir("csv_to_xlsx_files")
        # Si la carpeta csv_files existe, elimina todos los archivos CSV y luego elimina la carpeta csv_files
        if os.path.exists("csv_xlsx_files"):
            for xls_file_ in os.listdir("csv_xlsx_files"):
                a = open(os.path.join("csv_xlsx_files", xls_file_), "rb")
                a = a.close()

                os.remove(os.path.join("csv_xlsx_files", xls_file_))
            os.rmdir("csv_xlsx_files")
    deleteFiles()
    # Crea una carpeta para almacenar el PDF recibido
    os.mkdir("csv_to_xlsx_files")
    # Guardar el PDF recibido en el directorio doc_files
    file_name = os.path.join("csv_to_xlsx_files", file.filename)
    # Abre el archivo PDF en modo binario y lo guarda en el directorio doc_files
    with open(file_name, "wb") as buffer:
        # Lee el archivo PDF en trozos de 1024 bytes
        while chunk := await file.read(1024):
            # Escribe el trozo de bytes en el archivo PDF
            buffer.write(chunk)
    os.makedirs("csv_xlsx_files", exist_ok=True)

    # Creamos un objeto workbook de openpyxl
    workbook = Workbook()

    # Seleccionamos la hoja activa
    worksheet = workbook.active

    # Abrimos el archivo csv y lo leemos
    with open(file_name, 'r') as f:
        reader = csv.reader(f)

        # Iteramos sobre cada fila del archivo csv
        for row in reader:
            # Añadimos cada fila como una fila en la hoja de cálculo
            worksheet.append(row)

    # Guardamos el archivo xlsx
    workbook.save('csv_xlsx_files/archivo.xlsx')
    
    zip_file_name = "xlsx_files.zip"
    with zipfile.ZipFile(zip_file_name, "w") as zip_file:
        # Agrega todos los archivos DOC a la carpeta csv_xlsx_files
        for csv_file_name in os.listdir("csv_xlsx_files"):
            csv_file = os.path.join("csv_xlsx_files", csv_file_name)
            zip_file.write(csv_file, csv_file_name)
    # Mueve el archivo zip a la carpeta csv_files
    os.rename("xlsx_files.zip", "csv_xlsx_files/xlsx_files.zip")
    data = open("csv_xlsx_files/xlsx_files.zip", "rb")
    file_contents = data.read()
    data.close()
    response = StreamingResponse(BytesIO(file_contents), media_type="application/xlsx", headers={"Content-Disposition": "attachment; filename=csv_files.zip"})
    deleteFiles()
    return response

@app.post("/api/v1/convert-csv-to-xls")
async def convert_csv_xls(file: UploadFile = File(...)):
    import csv
    from openpyxl import Workbook

    def deleteFiles():
        # Si la carpeta doc_files existe, elimina todos los archivos PDF y luego elimina la carpeta doc_files
        if os.path.exists("csv_to_xls_files"):
            for csv_file_ in os.listdir("csv_to_xls_files"):
                os.remove(os.path.join("csv_to_xls_files", csv_file_))
            os.rmdir("csv_to_xls_files")
        # Si la carpeta csv_files existe, elimina todos los archivos CSV y luego elimina la carpeta csv_files
        if os.path.exists("csv_xls_files"):
            for xls_file_ in os.listdir("csv_xls_files"):
                a = open(os.path.join("csv_xls_files", xls_file_), "rb")
                a = a.close()

                os.remove(os.path.join("csv_xls_files", xls_file_))
            os.rmdir("csv_xls_files")
    deleteFiles()
    # Crea una carpeta para almacenar el PDF recibido
    os.mkdir("csv_to_xls_files")
    # Guardar el PDF recibido en el directorio doc_files
    file_name = os.path.join("csv_to_xls_files", file.filename)
    # Abre el archivo PDF en modo binario y lo guarda en el directorio doc_files
    with open(file_name, "wb") as buffer:
        # Lee el archivo PDF en trozos de 1024 bytes
        while chunk := await file.read(1024):
            # Escribe el trozo de bytes en el archivo PDF
            buffer.write(chunk)
    os.makedirs("csv_xls_files", exist_ok=True)

    # Creamos un objeto workbook de openpyxl
    workbook = Workbook()

    # Seleccionamos la hoja activa
    worksheet = workbook.active

    # Abrimos el archivo csv y lo leemos
    with open(file_name, 'r') as f:
        reader = csv.reader(f)

        # Iteramos sobre cada fila del archivo csv
        for row in reader:
            # Añadimos cada fila como una fila en la hoja de cálculo
            worksheet.append(row)

    # Guardamos el archivo xlsx
    workbook.save('csv_xls_files/archivo.xls')
    
    zip_file_name = "xls_files.zip"
    with zipfile.ZipFile(zip_file_name, "w") as zip_file:
        # Agrega todos los archivos DOC a la carpeta csv_xls_files
        for csv_file_name in os.listdir("csv_xls_files"):
            csv_file = os.path.join("csv_xls_files", csv_file_name)
            zip_file.write(csv_file, csv_file_name)
    # Mueve el archivo zip a la carpeta csv_files
    os.rename("xls_files.zip", "csv_xls_files/xls_files.zip")
    data = open("csv_xls_files/xls_files.zip", "rb")
    file_contents = data.read()
    data.close()
    response = StreamingResponse(BytesIO(file_contents), media_type="application/xls", headers={"Content-Disposition": "attachment; filename=csv_files.zip"})
    deleteFiles()
    return response

@app.post("/api/v1/convert-csv-to-docx")
async def convert_csv_docx(file: UploadFile = File(...)):
    import csv
    from docx import Document
    def deleteFiles():
        # Si la carpeta doc_files existe, elimina todos los archivos PDF y luego elimina la carpeta doc_files
        if os.path.exists("csv_to_docx_files"):
            for csv_file_ in os.listdir("csv_to_docx_files"):
                os.remove(os.path.join("csv_to_docx_files", csv_file_))
            os.rmdir("csv_to_docx_files")
        # Si la carpeta csv_files existe, elimina todos los archivos CSV y luego elimina la carpeta csv_files
        if os.path.exists("csv_docx_files"):
            for docx_file_ in os.listdir("csv_docx_files"):
                a = open(os.path.join("csv_docx_files", docx_file_), "rb")
                a = a.close()

                os.remove(os.path.join("csv_docx_files", docx_file_))
            os.rmdir("csv_docx_files")
    deleteFiles()
    # Crea una carpeta para almacenar el PDF recibido
    os.mkdir("csv_to_docx_files")
    # Guardar el PDF recibido en el directorio doc_files
    file_name = os.path.join("csv_to_docx_files", file.filename)
    # Abre el archivo PDF en modo binario y lo guarda en el directorio doc_files
    with open(file_name, "wb") as buffer:
        # Lee el archivo PDF en trozos de 1024 bytes
        while chunk := await file.read(1024):
            # Escribe el trozo de bytes en el archivo PDF
            buffer.write(chunk)
    os.makedirs("csv_docx_files", exist_ok=True)

    # Abre el archivo CSV con los datos
    with open(file_name, 'r') as archivo_csv:
        csv_reader = csv.reader(archivo_csv, delimiter=',')
        rows = list(csv_reader)
        encabezados = rows[0]

        # Crea un documento de Word
        document = Document()

        # Agrega una tabla al documento con la misma cantidad de columnas que el CSV
        table = document.add_table(rows=1, cols=len(encabezados))

        # Agrega los encabezados a la tabla
        hdr_cells = table.rows[0].cells
        for i in range(len(encabezados)):
            hdr_cells[i].text = encabezados[i]

        # Agrega los datos a la tabla
        for row in rows[1:]:
            row_cells = table.add_row().cells
            for i in range(len(encabezados)):
                row_cells[i].text = row[i]

        # Guarda el documento
        document.save('csv_docx_files/datos.docx')
    
    zip_file_name = "docx_files.zip"
    with zipfile.ZipFile(zip_file_name, "w") as zip_file:
        # Agrega todos los archivos DOC a la carpeta csv_docx_files
        for docx_file_name in os.listdir("csv_docx_files"):
            docx_file = os.path.join("csv_docx_files", docx_file_name)
            zip_file.write(docx_file, docx_file_name)
    # Mueve el archivo zip a la carpeta csv_files
    os.rename("docx_files.zip", "csv_docx_files/docx_files.zip")
    data = open("csv_docx_files/docx_files.zip", "rb")
    file_contents = data.read()
    data.close()
    response = StreamingResponse(BytesIO(file_contents), media_type="application/docx", headers={"Content-Disposition": "attachment; filename=csv_files.zip"})
    deleteFiles()
    return response

@app.post("/api/v1/convert-csv-to-doc")
async def convert_csv_doc(file: UploadFile = File(...)):
    import csv
    from docx import Document
    def deleteFiles():
        # Si la carpeta doc_files existe, elimina todos los archivos PDF y luego elimina la carpeta doc_files
        if os.path.exists("csv_to_doc_files"):
            for csv_file_ in os.listdir("csv_to_doc_files"):
                os.remove(os.path.join("csv_to_doc_files", csv_file_))
            os.rmdir("csv_to_doc_files")
        # Si la carpeta csv_files existe, elimina todos los archivos CSV y luego elimina la carpeta csv_files
        if os.path.exists("csv_doc_files"):
            for doc_file_ in os.listdir("csv_doc_files"):
                a = open(os.path.join("csv_doc_files", doc_file_), "rb")
                a = a.close()

                os.remove(os.path.join("csv_doc_files", doc_file_))
            os.rmdir("csv_doc_files")
    deleteFiles()
    # Crea una carpeta para almacenar el PDF recibido
    os.mkdir("csv_to_doc_files")
    # Guardar el PDF recibido en el directorio doc_files
    file_name = os.path.join("csv_to_doc_files", file.filename)
    # Abre el archivo PDF en modo binario y lo guarda en el directorio doc_files
    with open(file_name, "wb") as buffer:
        # Lee el archivo PDF en trozos de 1024 bytes
        while chunk := await file.read(1024):
            # Escribe el trozo de bytes en el archivo PDF
            buffer.write(chunk)
    os.makedirs("csv_doc_files", exist_ok=True)

    # Abre el archivo CSV con los datos
    with open(file_name, 'r') as archivo_csv:
        csv_reader = csv.reader(archivo_csv, delimiter=',')
        rows = list(csv_reader)
        encabezados = rows[0]

        # Crea un documento de Word
        document = Document()

        # Agrega una tabla al documento con la misma cantidad de columnas que el CSV
        table = document.add_table(rows=1, cols=len(encabezados))

        # Agrega los encabezados a la tabla
        hdr_cells = table.rows[0].cells
        for i in range(len(encabezados)):
            hdr_cells[i].text = encabezados[i]

        # Agrega los datos a la tabla
        for row in rows[1:]:
            row_cells = table.add_row().cells
            for i in range(len(encabezados)):
                row_cells[i].text = row[i]

        # Guarda el documento
        document.save('csv_doc_files/datos.doc')
    
    zip_file_name = "doc_files.zip"
    with zipfile.ZipFile(zip_file_name, "w") as zip_file:
        # Agrega todos los archivos DOC a la carpeta csv_doc_files
        for doc_file_name in os.listdir("csv_doc_files"):
            doc_file = os.path.join("csv_doc_files", doc_file_name)
            zip_file.write(doc_file, doc_file_name)
    # Mueve el archivo zip a la carpeta csv_files
    os.rename("doc_files.zip", "csv_doc_files/doc_files.zip")
    data = open("csv_doc_files/doc_files.zip", "rb")
    file_contents = data.read()
    data.close()
    response = StreamingResponse(BytesIO(file_contents), media_type="application/doc", headers={"Content-Disposition": "attachment; filename=csv_files.zip"})
    deleteFiles()
    return response


# Final conversor de CSV


# Inicio conversor de XLSX

@app.post("/api/v1/convert-xlsx-to-csv")
async def XLSX_CSV(file: UploadFile = File(...)):
    import pandas as pd
    def deleteFiles():
        # Si la carpeta pdf_files existe, elimina todos los archivos PDF y luego elimina la carpeta pdf_files
        if os.path.exists("XLSX_to_CSV_files"):
            for xlsx_file_ in os.listdir("XLSX_to_CSV_files"):
                os.remove(os.path.join("XLSX_to_CSV_files", xlsx_file_))
            os.rmdir("XLSX_to_CSV_files")
        # Si la carpeta csv_files existe, elimina todos los archivos CSV y luego elimina la carpeta csv_files
        if os.path.exists("XLSX_CSV_files"):
            for csv_file_ in os.listdir("XLSX_CSV_files"):
                a = open(os.path.join("XLSX_CSV_files", csv_file_), "rb")
                a = a.close()

                os.remove(os.path.join("XLSX_CSV_files", csv_file_))
            os.rmdir("XLSX_CSV_files")
    deleteFiles()
    # Crea una carpeta para almacenar el PDF recibido
    os.mkdir("XLSX_to_CSV_files")
    # Guardar el PDF recibido en el directorio pdf_files
    file_name = os.path.join("XLSX_to_CSV_files", file.filename)
    # Abre el archivo PDF en modo binario y lo guarda en el directorio pdf_files
    with open(file_name, "wb") as buffer:
        # Lee el archivo PDF en trozos de 1024 bytes
        while chunk := await file.read(1024):
            # Escribe el trozo de bytes en el archivo PDF
            buffer.write(chunk)
    os.makedirs("XLSX_CSV_files", exist_ok=True)


    # Cargamos el archivo xlsx en un DataFrame de pandas
    df = pd.read_excel(file_name)

    # Guardamos el DataFrame como archivo csv
    df.to_csv('XLSX_CSV_files/converted.csv', index=False)

    zip_file_name = "csv_file.zip"
    with zipfile.ZipFile(zip_file_name, "w") as zip_file:
        # Agrega todos los archivos DOC a la carpeta XLSX_CSV_files
        for pdf_file_name in os.listdir("XLSX_CSV_files"):
            csv_file = os.path.join("XLSX_CSV_files", pdf_file_name)
            zip_file.write(csv_file, pdf_file_name)
    # Mueve el archivo zip a la carpeta csv_files
    os.rename("csv_file.zip", "XLSX_CSV_files/csv_file.zip")
    data = open("XLSX_CSV_files/csv_file.zip", "rb")
    file_contents = data.read()
    data.close()
    response = StreamingResponse(BytesIO(file_contents), media_type="application/csv", headers={"Content-Disposition": "attachment; filename=csv_files.zip"})
    deleteFiles()
    return response

@app.post("/api/v1/convert-xlsx-to-pdf")
async def convert_xlsx_pdf(file: UploadFile = File(...)):
    import pandas as pd
    from fpdf import FPDF
    def deleteFiles():
        # Si la carpeta pdf_files existe, elimina todos los archivos PDF y luego elimina la carpeta pdf_files
        if os.path.exists("XLSX_to_PDF_files"):
            for xlsx_file_ in os.listdir("XLSX_to_PDF_files"):
                os.remove(os.path.join("XLSX_to_PDF_files", xlsx_file_))
            os.rmdir("XLSX_to_PDF_files")
        # Si la carpeta csv_files existe, elimina todos los archivos CSV y luego elimina la carpeta csv_files
        if os.path.exists("XLSX_PDF_files"):
            for pdf_file_ in os.listdir("XLSX_PDF_files"):
                a = open(os.path.join("XLSX_PDF_files", pdf_file_), "rb")
                a = a.close()

                os.remove(os.path.join("XLSX_PDF_files", pdf_file_))
            os.rmdir("XLSX_PDF_files")
    deleteFiles()
    # Crea una carpeta para almacenar el PDF recibido
    os.mkdir("XLSX_to_PDF_files")
    # Guardar el PDF recibido en el directorio pdf_files
    file_name = os.path.join("XLSX_to_PDF_files", file.filename)
    # Abre el archivo PDF en modo binario y lo guarda en el directorio pdf_files
    with open(file_name, "wb") as buffer:
        # Lee el archivo PDF en trozos de 1024 bytes
        while chunk := await file.read(1024):
            # Escribe el trozo de bytes en el archivo PDF
            buffer.write(chunk)
    os.makedirs("XLSX_PDF_files", exist_ok=True)

    def convert_xls_to_pdf(xls_file_path, pdf_file_path):
        # Leer el archivo XLS con pandas y reemplazar los NaN por una cadena vacía
        df = pd.read_excel(xls_file_path).fillna("")

        # Crear un objeto FPDF
        pdf = FPDF()

        # Agregar una página al PDF
        pdf.add_page()

        # Establecer la fuente y el tamaño de letra
        pdf.set_font("Arial", size=12)

        # Recorrer las filas y columnas del DataFrame y agregarlas al PDF
        for index, row in df.iterrows():
            for col, value in enumerate(row):
                pdf.cell(40, 10, str(value), border=1)
            pdf.ln()

        # Guardar el archivo PDF
        pdf.output(pdf_file_path)

    convert_xls_to_pdf(file_name, "XLSX_PDF_files/converted.pdf")
    
    zip_file_name = "format_file.zip"
    with zipfile.ZipFile(zip_file_name, "w") as zip_file:
        # Agrega todos los archivos DOC a la carpeta XLSX_PDF_files
        for pdf_file_name in os.listdir("XLSX_PDF_files"):
            csv_file = os.path.join("XLSX_PDF_files", pdf_file_name)
            zip_file.write(csv_file, pdf_file_name)
    # Mueve el archivo zip a la carpeta csv_files
    os.rename("format_file.zip", "XLSX_PDF_files/format_file.zip")
    data = open("XLSX_PDF_files/format_file.zip", "rb")
    file_contents = data.read()
    data.close()
    response = StreamingResponse(BytesIO(file_contents), media_type="application/pdf", headers={"Content-Disposition": "attachment; filename=csv_files.zip"})
    deleteFiles()
    return response

@app.post("/api/v1/convert-xlsx-to-xls")
async def convert_xlsx_xls(file: UploadFile = File(...)):
    from openpyxl import load_workbook
    from openpyxl.workbook import Workbook
    def deleteFiles():
        # Si la carpeta pdf_files existe, elimina todos los archivos PDF y luego elimina la carpeta pdf_files
        if os.path.exists("XLSX_to_XLS_files"):
            for xlsx_file_ in os.listdir("XLSX_to_XLS_files"):
                os.remove(os.path.join("XLSX_to_XLS_files", xlsx_file_))
            os.rmdir("XLSX_to_XLS_files")
        # Si la carpeta csv_files existe, elimina todos los archivos CSV y luego elimina la carpeta csv_files
        if os.path.exists("XLSX_XLS_files"):
            for xls_file_ in os.listdir("XLSX_XLS_files"):
                a = open(os.path.join("XLSX_XLS_files", xls_file_), "rb")
                a = a.close()

                os.remove(os.path.join("XLSX_XLS_files", xls_file_))
            os.rmdir("XLSX_XLS_files")
    deleteFiles()
    # Crea una carpeta para almacenar el PDF recibido
    os.mkdir("XLSX_to_XLS_files")
    # Guardar el PDF recibido en el directorio pdf_files
    file_name = os.path.join("XLSX_to_XLS_files", file.filename)
    # Abre el archivo PDF en modo binario y lo guarda en el directorio pdf_files
    with open(file_name, "wb") as buffer:
        # Lee el archivo PDF en trozos de 1024 bytes
        while chunk := await file.read(1024):
            # Escribe el trozo de bytes en el archivo PDF
            buffer.write(chunk)
    os.makedirs("XLSX_XLS_files", exist_ok=True)


    # Cargar el archivo xlsx
    wb = load_workbook(file_name)

    # Crear un nuevo libro de trabajo en formato xls
    new_wb = Workbook()

    # Copiar cada hoja del archivo xlsx al nuevo libro de trabajo xls
    for sheet in wb:
        new_sheet = new_wb.create_sheet(title=sheet.title)
        for row in sheet.iter_rows(values_only=True):
            new_sheet.append(row)

    # Guardar el nuevo libro de trabajo en formato xls
    new_wb.save("XLSX_XLS_files/converted.xls")

    zip_file_name = "xls_file.zip"
    with zipfile.ZipFile(zip_file_name, "w") as zip_file:
        # Agrega todos los archivos DOC a la carpeta XLSX_XLS_files
        for pdf_file_name in os.listdir("XLSX_XLS_files"):
            csv_file = os.path.join("XLSX_XLS_files", pdf_file_name)
            zip_file.write(csv_file, pdf_file_name)
    # Mueve el archivo zip a la carpeta csv_files
    os.rename("xls_file.zip", "XLSX_XLS_files/xls_file.zip")
    data = open("XLSX_XLS_files/xls_file.zip", "rb")
    file_contents = data.read()
    data.close()
    response = StreamingResponse(BytesIO(file_contents), media_type="application/xls", headers={"Content-Disposition": "attachment; filename=csv_files.zip"})
    deleteFiles()
    return response 

@app.post("/api/v1/convert-xlsx-to-doc")
async def convert_xlsx_doc(file: UploadFile = File(...)):
    import openpyxl
    import docx
    def deleteFiles():
        # Si la carpeta pdf_files existe, elimina todos los archivos PDF y luego elimina la carpeta pdf_files
        if os.path.exists("XLSX_to_DOC_files"):
            for xls_file_ in os.listdir("XLSX_to_DOC_files"):
                os.remove(os.path.join("XLSX_to_DOC_files", xls_file_))
            os.rmdir("XLSX_to_DOC_files")
        # Si la carpeta doc_files existe, elimina todos los archivos CSV y luego elimina la carpeta doc_files
        if os.path.exists("XLSX_DOC_files"):
            for doc_file_ in os.listdir("XLSX_DOC_files"):
                a = open(os.path.join("XLSX_DOC_files", doc_file_), "rb")
                a = a.close()

                os.remove(os.path.join("XLSX_DOC_files", doc_file_))
            os.rmdir("XLSX_DOC_files")
    deleteFiles()
    # Crea una carpeta para almacenar el PDF recibido
    os.mkdir("XLSX_to_DOC_files")
    # Guardar el PDF recibido en el directorio pdf_files
    file_name = os.path.join("XLSX_to_DOC_files", file.filename)
    # Abre el archivo PDF en modo binario y lo guarda en el directorio pdf_files
    with open(file_name, "wb") as buffer:
        # Lee el archivo PDF en trozos de 1024 bytes
        while chunk := await file.read(1024):
            # Escribe el trozo de bytes en el archivo PDF
            buffer.write(chunk)
    os.makedirs("XLSX_DOC_files", exist_ok=True)

    # Abrimos el archivo de Excel
    workbook = openpyxl.load_workbook(file_name)

    # Obtenemos la primera hoja del libro
    worksheet = workbook.active

    # Creamos un nuevo archivo de Word
    document = docx.Document()

    # Obtenemos los datos de cada fila y los agregamos al documento de Word
    for row in worksheet.iter_rows():
        cells = [str(cell.value) for cell in row]
        document.add_paragraph(' | '.join(cells))

    # Guardamos el archivo de Word
    document.save('XLSX_DOC_files/converted.doc')
    
    zip_file_name = "doc_file.zip"
    with zipfile.ZipFile(zip_file_name, "w") as zip_file:
        # Agrega todos los archivos DOC a la carpeta XLSX_DOC_files
        for doc_file_name in os.listdir("XLSX_DOC_files"):
            doc_file = os.path.join("XLSX_DOC_files", doc_file_name)
            zip_file.write(doc_file, doc_file_name)
    # Mueve el archivo zip a la carpeta doc_files
    os.rename("doc_file.zip", "XLSX_DOC_files/doc_file.zip")
    data = open("XLSX_DOC_files/doc_file.zip", "rb")
    file_contents = data.read()
    data.close()
    response = StreamingResponse(BytesIO(file_contents), media_type="application/doc", headers={"Content-Disposition": "attachment; filename=doc_files.zip"})
    deleteFiles()
    return response

@app.post("/api/v1/convert-xlsx-to-docx")
async def convert_xlsx_docx(file: UploadFile = File(...)):
    import openpyxl
    import docx
    def deleteFiles():
        # Si la carpeta pdf_files existe, elimina todos los archivos PDF y luego elimina la carpeta pdf_files
        if os.path.exists("XLSX_to_DOCX_files"):
            for xls_file_ in os.listdir("XLSX_to_DOCX_files"):
                os.remove(os.path.join("XLSX_to_DOCX_files", xls_file_))
            os.rmdir("XLSX_to_DOCX_files")
        # Si la carpeta docX_files existe, elimina todos los archivos CSV y luego elimina la carpeta docX_files
        if os.path.exists("XLSX_DOCX_files"):
            for docx_file_ in os.listdir("XLSX_DOCX_files"):
                a = open(os.path.join("XLSX_DOCX_files", docx_file_), "rb")
                a = a.close()

                os.remove(os.path.join("XLSX_DOCX_files", docx_file_))
            os.rmdir("XLSX_DOCX_files")
    deleteFiles()
    # Crea una carpeta para almacenar el PDF recibido
    os.mkdir("XLSX_to_DOCX_files")
    # Guardar el PDF recibido en el directorio pdf_files
    file_name = os.path.join("XLSX_to_DOCX_files", file.filename)
    # Abre el archivo PDF en modo binario y lo guarda en el directorio pdf_files
    with open(file_name, "wb") as buffer:
        # Lee el archivo PDF en trozos de 1024 bytes
        while chunk := await file.read(1024):
            # Escribe el trozo de bytes en el archivo PDF
            buffer.write(chunk)
    os.makedirs("XLSX_DOCX_files", exist_ok=True)

    # Abrimos el archivo de Excel
    workbook = openpyxl.load_workbook(file_name)

    # Obtenemos la primera hoja del libro
    worksheet = workbook.active

    # Creamos un nuevo archivo de Word
    document = docx.Document()

    # Obtenemos los datos de cada fila y los agregamos al documento de Word
    for row in worksheet.iter_rows():
        cells = [str(cell.value) for cell in row]
        document.add_paragraph(' | '.join(cells))

    # Guardamos el archivo de Word
    document.save('XLSX_DOCX_files/converted.docx')
    
    zip_file_name = "docx_file.zip"
    with zipfile.ZipFile(zip_file_name, "w") as zip_file:
        # Agrega todos los archivos DOCX a la carpeta XLSX_DOCX_files
        for docx_file_name in os.listdir("XLSX_DOCX_files"):
            docx_file = os.path.join("XLSX_DOCX_files", docx_file_name)
            zip_file.write(docx_file, docx_file_name)
    # Mueve el archivo zip a la carpeta docx_files
    os.rename("docx_file.zip", "XLSX_DOCX_files/docx_file.zip")
    data = open("XLSX_DOCX_files/docx_file.zip", "rb")
    file_contents = data.read()
    data.close()
    response = StreamingResponse(BytesIO(file_contents), media_type="application/docx", headers={"Content-Disposition": "attachment; filename=docX_files.zip"})
    deleteFiles()
    return response


# Final conversor de XLSX


# Inicio converor de XLS

@app.post("/api/v1/convert-xls-csv")
async def convert_xls_csv(file: UploadFile = File(...)):
    import pandas as pd
    def deleteFiles():
        # Si la carpeta pdf_files existe, elimina todos los archivos PDF y luego elimina la carpeta pdf_files
        if os.path.exists("XLS_to_CSV_files"):
            for xls_file_ in os.listdir("XLS_to_CSV_files"):
                os.remove(os.path.join("XLS_to_CSV_files", xls_file_))
            os.rmdir("XLS_to_CSV_files")
        # Si la carpeta csv_files existe, elimina todos los archivos CSV y luego elimina la carpeta csv_files
        if os.path.exists("XLS_CSV_files"):
            for csv_file_ in os.listdir("XLS_CSV_files"):
                a = open(os.path.join("XLS_CSV_files", csv_file_), "rb")
                a = a.close()

                os.remove(os.path.join("XLS_CSV_files", csv_file_))
            os.rmdir("XLS_CSV_files")
    deleteFiles()
    # Crea una carpeta para almacenar el PDF recibido
    os.mkdir("XLS_to_CSV_files")
    # Guardar el PDF recibido en el directorio pdf_files
    file_name = os.path.join("XLS_to_CSV_files", file.filename)
    # Abre el archivo PDF en modo binario y lo guarda en el directorio pdf_files
    with open(file_name, "wb") as buffer:
        # Lee el archivo PDF en trozos de 1024 bytes
        while chunk := await file.read(1024):
            # Escribe el trozo de bytes en el archivo PDF
            buffer.write(chunk)
    os.makedirs("XLS_CSV_files", exist_ok=True)


    # Cargamos el archivo xls en un DataFrame de pandas
    df = pd.read_excel(file_name)

    # Guardamos el DataFrame como archivo csv
    df.to_csv('XLS_CSV_files/converted.csv', index=False)

    zip_file_name = "csv_file.zip"
    with zipfile.ZipFile(zip_file_name, "w") as zip_file:
        # Agrega todos los archivos DOC a la carpeta XLS_CSV_files
        for pdf_file_name in os.listdir("XLS_CSV_files"):
            csv_file = os.path.join("XLS_CSV_files", pdf_file_name)
            zip_file.write(csv_file, pdf_file_name)
    # Mueve el archivo zip a la carpeta csv_files
    os.rename("csv_file.zip", "XLS_CSV_files/csv_file.zip")
    data = open("XLS_CSV_files/csv_file.zip", "rb")
    file_contents = data.read()
    data.close()
    response = StreamingResponse(BytesIO(file_contents), media_type="application/csv", headers={"Content-Disposition": "attachment; filename=csv_files.zip"})
    deleteFiles()
    return response

@app.post("/api/v1/convert-xls-to-pdf")
async def convert_xls_pdf(file: UploadFile = File(...)):
    import pandas as pd
    from fpdf import FPDF
    def deleteFiles():
        # Si la carpeta pdf_files existe, elimina todos los archivos PDF y luego elimina la carpeta pdf_files
        if os.path.exists("XLS_to_PDF_files"):
            for xls_file_ in os.listdir("XLS_to_PDF_files"):
                os.remove(os.path.join("XLS_to_PDF_files", xls_file_))
            os.rmdir("XLS_to_PDF_files")
        # Si la carpeta csv_files existe, elimina todos los archivos CSV y luego elimina la carpeta csv_files
        if os.path.exists("XLS_PDF_files"):
            for pdf_file_ in os.listdir("XLS_PDF_files"):
                a = open(os.path.join("XLS_PDF_files", pdf_file_), "rb")
                a = a.close()

                os.remove(os.path.join("XLS_PDF_files", pdf_file_))
            os.rmdir("XLS_PDF_files")
    deleteFiles()
    # Crea una carpeta para almacenar el PDF recibido
    os.mkdir("XLS_to_PDF_files")
    # Guardar el PDF recibido en el directorio pdf_files
    file_name = os.path.join("XLS_to_PDF_files", file.filename)
    # Abre el archivo PDF en modo binario y lo guarda en el directorio pdf_files
    with open(file_name, "wb") as buffer:
        # Lee el archivo PDF en trozos de 1024 bytes
        while chunk := await file.read(1024):
            # Escribe el trozo de bytes en el archivo PDF
            buffer.write(chunk)
    os.makedirs("XLS_PDF_files", exist_ok=True)

    def convert_xls_to_pdf(xls_file_path, pdf_file_path):
        # Leer el archivo XLS con pandas y reemplazar los NaN por una cadena vacía
        df = pd.read_excel(xls_file_path).fillna("")

        # Crear un objeto FPDF
        pdf = FPDF()

        # Agregar una página al PDF
        pdf.add_page()

        # Establecer la fuente y el tamaño de letra
        pdf.set_font("Arial", size=12)

        # Recorrer las filas y columnas del DataFrame y agregarlas al PDF
        for index, row in df.iterrows():
            for col, value in enumerate(row):
                pdf.cell(40, 10, str(value), border=1)
            pdf.ln()

        # Guardar el archivo PDF
        pdf.output(pdf_file_path)

    convert_xls_to_pdf(file_name, "XLS_PDF_files/converted.pdf")
    
    zip_file_name = "format_file.zip"
    with zipfile.ZipFile(zip_file_name, "w") as zip_file:
        # Agrega todos los archivos DOC a la carpeta XLS_PDF_files
        for pdf_file_name in os.listdir("XLS_PDF_files"):
            csv_file = os.path.join("XLS_PDF_files", pdf_file_name)
            zip_file.write(csv_file, pdf_file_name)
    # Mueve el archivo zip a la carpeta csv_files
    os.rename("format_file.zip", "XLS_PDF_files/format_file.zip")
    data = open("XLS_PDF_files/format_file.zip", "rb")
    file_contents = data.read()
    data.close()
    response = StreamingResponse(BytesIO(file_contents), media_type="application/format", headers={"Content-Disposition": "attachment; filename=csv_files.zip"})
    deleteFiles()
    return response

@app.post("/api/v1/convert-xls-to-doc")
async def convert_xls_doc(file: UploadFile = File(...)):
    import xlrd
    import docx
    def deleteFiles():
        # Si la carpeta pdf_files existe, elimina todos los archivos PDF y luego elimina la carpeta pdf_files
        if os.path.exists("XLS_to_DOC_files"):
            for xls_file_ in os.listdir("XLS_to_DOC_files"):
                os.remove(os.path.join("XLS_to_DOC_files", xls_file_))
            os.rmdir("XLS_to_DOC_files")
        # Si la carpeta csv_files existe, elimina todos los archivos CSV y luego elimina la carpeta csv_files
        if os.path.exists("XLS_DOC_files"):
            for doc_file_ in os.listdir("XLS_DOC_files"):
                a = open(os.path.join("XLS_DOC_files", doc_file_), "rb")
                a = a.close()

                os.remove(os.path.join("XLS_DOC_files", doc_file_))
            os.rmdir("XLS_DOC_files")
    deleteFiles()
    # Crea una carpeta para almacenar el PDF recibido
    os.mkdir("XLS_to_DOC_files")
    # Guardar el PDF recibido en el directorio pdf_files
    file_name = os.path.join("XLS_to_DOC_files", file.filename)
    # Abre el archivo PDF en modo binario y lo guarda en el directorio pdf_files
    with open(file_name, "wb") as buffer:
        # Lee el archivo PDF en trozos de 1024 bytes
        while chunk := await file.read(1024):
            # Escribe el trozo de bytes en el archivo PDF
            buffer.write(chunk)
    os.makedirs("XLS_DOC_files", exist_ok=True)


    word_file = 'XLS_DOC_files/archivo.doc'

    # Abrir el archivo de Excel
    workbook = xlrd.open_workbook(file_name)

    # Seleccionar la primera hoja
    worksheet = workbook.sheet_by_index(0)

    # Obtener el número de filas y columnas
    num_rows = worksheet.nrows
    num_cols = worksheet.ncols

    # Crear un nuevo documento de Word
    document = docx.Document()

    # Añadir una tabla vacía con el mismo número de filas y columnas que la hoja de Excel
    table = document.add_table(num_rows, num_cols)

    # Añadir el contenido de la hoja de Excel a la tabla de Word
    for i in range(num_rows):
        for j in range(num_cols):
            table.cell(i, j).text = str(worksheet.cell_value(i, j))

    # Guardar el documento de Word
    document.save(word_file)
    
    zip_file_name = "format_file.zip"
    with zipfile.ZipFile(zip_file_name, "w") as zip_file:
        # Agrega todos los archivos DOC a la carpeta XLS_DOC_files
        for pdf_file_name in os.listdir("XLS_DOC_files"):
            csv_file = os.path.join("XLS_DOC_files", pdf_file_name)
            zip_file.write(csv_file, pdf_file_name)
    # Mueve el archivo zip a la carpeta csv_files
    os.rename("format_file.zip", "XLS_DOC_files/format_file.zip")
    data = open("XLS_DOC_files/format_file.zip", "rb")
    file_contents = data.read()
    data.close()
    response = StreamingResponse(BytesIO(file_contents), media_type="application/format", headers={"Content-Disposition": "attachment; filename=csv_files.zip"})
    deleteFiles()
    return response

@app.post("/api/v1/convert-xls-to-docx")
async def convert_xls_docx(file: UploadFile = File(...)):
    import xlrd
    import docx
    def deleteFiles():
        # Si la carpeta pdf_files existe, elimina todos los archivos PDF y luego elimina la carpeta pdf_files
        if os.path.exists("XLS_to_DOCX_files"):
            for xls_file_ in os.listdir("XLS_to_DOCX_files"):
                os.remove(os.path.join("XLS_to_DOCX_files", xls_file_))
            os.rmdir("XLS_to_DOCX_files")
        # Si la carpeta csv_files existe, elimina todos los archivos CSV y luego elimina la carpeta csv_files
        if os.path.exists("XLS_DOCX_files"):
            for docx_file_ in os.listdir("XLS_DOCX_files"):
                a = open(os.path.join("XLS_DOCX_files", docx_file_), "rb")
                a = a.close()

                os.remove(os.path.join("XLS_DOCX_files", docx_file_))
            os.rmdir("XLS_DOCX_files")
    deleteFiles()
    # Crea una carpeta para almacenar el PDF recibido
    os.mkdir("XLS_to_DOCX_files")
    # Guardar el PDF recibido en el directorio pdf_files
    file_name = os.path.join("XLS_to_DOCX_files", file.filename)
    # Abre el archivo PDF en modo binario y lo guarda en el directorio pdf_files
    with open(file_name, "wb") as buffer:
        # Lee el archivo PDF en trozos de 1024 bytes
        while chunk := await file.read(1024):
            # Escribe el trozo de bytes en el archivo PDF
            buffer.write(chunk)
    os.makedirs("XLS_DOCX_files", exist_ok=True)


    word_file = 'XLS_DOCX_files/archivo.docx'

    # Abrir el archivo de Excel
    workbook = xlrd.open_workbook(file_name)

    # Seleccionar la primera hoja
    worksheet = workbook.sheet_by_index(0)

    # Obtener el número de filas y columnas
    num_rows = worksheet.nrows
    num_cols = worksheet.ncols

    # Crear un nuevo documento de Word
    document = docx.Document()

    # Añadir una tabla vacía con el mismo número de filas y columnas que la hoja de Excel
    table = document.add_table(num_rows, num_cols)

    # Añadir el contenido de la hoja de Excel a la tabla de Word
    for i in range(num_rows):
        for j in range(num_cols):
            table.cell(i, j).text = str(worksheet.cell_value(i, j))

    # Guardar el documento de Word
    document.save(word_file)
    
    zip_file_name = "docx_file.zip"
    with zipfile.ZipFile(zip_file_name, "w") as zip_file:
        # Agrega todos los archivos DOC a la carpeta XLS_DOCX_files
        for pdf_file_name in os.listdir("XLS_DOCX_files"):
            csv_file = os.path.join("XLS_DOCX_files", pdf_file_name)
            zip_file.write(csv_file, pdf_file_name)
    # Mueve el archivo zip a la carpeta csv_files
    os.rename("docx_file.zip", "XLS_DOCX_files/docx_file.zip")
    data = open("XLS_DOCX_files/docx_file.zip", "rb")
    file_contents = data.read()
    data.close()
    response = StreamingResponse(BytesIO(file_contents), media_type="application/docx", headers={"Content-Disposition": "attachment; filename=csv_files.zip"})
    deleteFiles()
    return response

@app.post("/api/v1/convert-xls-xlsx")
async def convert_xls_xlsx(file: UploadFile = File(...)):
    import openpyxl
    import xlwt
    def deleteFiles():
        # Si la carpeta pdf_files existe, elimina todos los archivos PDF y luego elimina la carpeta pdf_files
        if os.path.exists("XLS_to_XLSX_files"):
            for xls_file_ in os.listdir("XLS_to_XLSX_files"):
                os.remove(os.path.join("XLS_to_XLSX_files", xls_file_))
            os.rmdir("XLS_to_XLSX_files")
        # Si la carpeta csv_files existe, elimina todos los archivos CSV y luego elimina la carpeta csv_files
        if os.path.exists("XLS_XLSX_files"):
            for xlsx_file_ in os.listdir("XLS_XLSX_files"):
                a = open(os.path.join("XLS_XLSX_files", xlsx_file_), "rb")
                a = a.close()

                os.remove(os.path.join("XLS_XLSX_files", xlsx_file_))
            os.rmdir("XLS_XLSX_files")
    deleteFiles()
    # Crea una carpeta para almacenar el PDF recibido
    os.mkdir("XLS_to_XLSX_files")
    # Guardar el PDF recibido en el directorio pdf_files
    file_name = os.path.join("XLS_to_XLSX_files", file.filename)
    # Abre el archivo PDF en modo binario y lo guarda en el directorio pdf_files
    with open(file_name, "wb") as buffer:
        # Lee el archivo PDF en trozos de 1024 bytes
        while chunk := await file.read(1024):
            # Escribe el trozo de bytes en el archivo PDF
            buffer.write(chunk)
    os.makedirs("XLS_XLSX_files", exist_ok=True)

    # Abre el archivo xlsx
    wb = openpyxl.load_workbook(file_name)

    # Crea un nuevo libro de trabajo xls
    new_wb = xlwt.Workbook()

    # Copia todas las hojas del libro de trabajo xlsx al libro de trabajo xls
    for sheetname in wb.sheetnames:
        ws = wb[sheetname]
        new_ws = new_wb.add_sheet(sheetname)
        for row in ws.rows:
            for col, cell in enumerate(row):
                new_ws.write(cell.row-1, cell.column-1, cell.value)

    # Guarda el archivo xls
    new_wb.save('archivo.xls')

    zip_file_name = "format_file.zip"
    with zipfile.ZipFile(zip_file_name, "w") as zip_file:
        # Agrega todos los archivos DOC a la carpeta XLS_XLSX_files
        for pdf_file_name in os.listdir("XLS_XLSX_files"):
            csv_file = os.path.join("XLS_XLSX_files", pdf_file_name)
            zip_file.write(csv_file, pdf_file_name)
    # Mueve el archivo zip a la carpeta csv_files
    os.rename("format_file.zip", "XLS_XLSX_files/format_file.zip")
    data = open("XLS_XLSX_files/format_file.zip", "rb")
    file_contents = data.read()
    data.close()
    response = StreamingResponse(BytesIO(file_contents), media_type="application/xlsx", headers={"Content-Disposition": "attachment; filename=csv_files.zip"})
    deleteFiles()
    return response


# Final conversor de XLS


# Inicio conversor de DOC

@app.post("/api/v1/convert-doc-to-csv")
async def convert_doc_csv(file: UploadFile = File(...)):
    import pandas as pd
    import docx
    def deleteFiles():
        # Si la carpeta pdf_files existe, elimina todos los archivos PDF y luego elimina la carpeta pdf_files
        if os.path.exists("doc_to_csv_files"):
            for doc_file_ in os.listdir("doc_to_csv_files"):
                os.remove(os.path.join("doc_to_csv_files", doc_file_))
            os.rmdir("doc_to_csv_files")
        # Si la carpeta csv_files existe, elimina todos los archivos CSV y luego elimina la carpeta csv_files
        if os.path.exists("doc_csv_files"):
            for csv_file_ in os.listdir("doc_csv_files"):
                a = open(os.path.join("doc_csv_files", csv_file_), "rb")
                a = a.close()

                os.remove(os.path.join("doc_csv_files", csv_file_))
            os.rmdir("doc_csv_files")
    deleteFiles()
    # Crea una carpeta para almacenar el PDF recibido
    os.mkdir("doc_to_csv_files")
    # Guardar el PDF recibido en el directorio pdf_files
    file_name = os.path.join("doc_to_csv_files", file.filename)
    # Abre el archivo PDF en modo binario y lo guarda en el directorio pdf_files
    with open(file_name, "wb") as buffer:
        # Lee el archivo PDF en trozos de 1024 bytes
        while chunk := await file.read(1024):
            # Escribe el trozo de bytes en el archivo PDF
            buffer.write(chunk)
    os.makedirs("doc_csv_files", exist_ok=True)


    # Cargar el archivo de Word
    doc = docx.Document(file_name)

    # Inicializar una lista vacía para almacenar los datos de la tabla
    table_data = []

    # Iterar por cada tabla en el documento
    for table in doc.tables:
        # Iterar por cada fila de la tabla
        for row in table.rows:
            # Inicializar una lista vacía para almacenar los datos de la fila
            row_data = []
            # Iterar por cada celda de la fila
            for cell in row.cells:
                # Agregar el contenido de la celda a la lista de datos de la fila
                row_data.append(cell.text)
            # Agregar los datos de la fila a la lista de datos de la tabla
            table_data.append(row_data)

    # Convertir los datos de la tabla en un DataFrame de Pandas
    df = pd.DataFrame(table_data)

    # Guardar el DataFrame en un archivo de Excel
    df.to_csv("doc_csv_files/converted.csv", index=False)

    zip_file_name = "csv_file.zip"
    with zipfile.ZipFile(zip_file_name, "w") as zip_file:
        # Agrega todos los archivos DOC a la carpeta doc_csv_files
        for pdf_file_name in os.listdir("doc_csv_files"):
            csv_file = os.path.join("doc_csv_files", pdf_file_name)
            zip_file.write(csv_file, pdf_file_name)
    # Mueve el archivo zip a la carpeta csv_files
    os.rename("csv_file.zip", "doc_csv_files/csv_file.zip")
    data = open("doc_csv_files/csv_file.zip", "rb")
    file_contents = data.read()
    data.close()
    response = StreamingResponse(BytesIO(file_contents), media_type="application/csv", headers={"Content-Disposition": "attachment; filename=csv_files.zip"})
    deleteFiles()
    return response

@app.post("/api/v1/convert-doc-to-xlsx")
async def convert_doc_xlsx(file: UploadFile = File(...)):
    import pandas as pd
    import docx
    def deleteFiles():
        # Si la carpeta pdf_files existe, elimina todos los archivos PDF y luego elimina la carpeta pdf_files
        if os.path.exists("doc_to_xlsx_files"):
            for doc_file_ in os.listdir("doc_to_xlsx_files"):
                os.remove(os.path.join("doc_to_xlsx_files", doc_file_))
            os.rmdir("doc_to_xlsx_files")
        # Si la carpeta csv_files existe, elimina todos los archivos CSV y luego elimina la carpeta csv_files
        if os.path.exists("doc_xlsx_files"):
            for xlsx_file_ in os.listdir("doc_xlsx_files"):
                a = open(os.path.join("doc_xlsx_files", xlsx_file_), "rb")
                a = a.close()

                os.remove(os.path.join("doc_xlsx_files", xlsx_file_))
            os.rmdir("doc_xlsx_files")
    deleteFiles()
    # Crea una carpeta para almacenar el PDF recibido
    os.mkdir("doc_to_xlsx_files")
    # Guardar el PDF recibido en el directorio pdf_files
    file_name = os.path.join("doc_to_xlsx_files", file.filename)
    # Abre el archivo PDF en modo binario y lo guarda en el directorio pdf_files
    with open(file_name, "wb") as buffer:
        # Lee el archivo PDF en trozos de 1024 bytes
        while chunk := await file.read(1024):
            # Escribe el trozo de bytes en el archivo PDF
            buffer.write(chunk)
    os.makedirs("doc_xlsx_files", exist_ok=True)


    # Cargar el archivo de Word
    doc = docx.Document(file_name)

    # Inicializar una lista vacía para almacenar los datos de la tabla
    table_data = []

    # Iterar por cada tabla en el documento
    for table in doc.tables:
        # Iterar por cada fila de la tabla
        for row in table.rows:
            # Inicializar una lista vacía para almacenar los datos de la fila
            row_data = []
            # Iterar por cada celda de la fila
            for cell in row.cells:
                # Agregar el contenido de la celda a la lista de datos de la fila
                row_data.append(cell.text)
            # Agregar los datos de la fila a la lista de datos de la tabla
            table_data.append(row_data)

    # Convertir los datos de la tabla en un DataFrame de Pandas
    df = pd.DataFrame(table_data)

    # Guardar el DataFrame en un archivo de Excel
    df.to_excel("doc_xlsx_files/converted.xlsx", index=False)

    zip_file_name = "xlsx_file.zip"
    with zipfile.ZipFile(zip_file_name, "w") as zip_file:
        # Agrega todos los archivos DOC a la carpeta doc_xlsx_files
        for pdf_file_name in os.listdir("doc_xlsx_files"):
            csv_file = os.path.join("doc_xlsx_files", pdf_file_name)
            zip_file.write(csv_file, pdf_file_name)
    # Mueve el archivo zip a la carpeta csv_files
    os.rename("xlsx_file.zip", "doc_xlsx_files/xlsx_file.zip")
    data = open("doc_xlsx_files/xlsx_file.zip", "rb")
    file_contents = data.read()
    data.close()
    response = StreamingResponse(BytesIO(file_contents), media_type="application/xlsx", headers={"Content-Disposition": "attachment; filename=csv_files.zip"})
    deleteFiles()
    return response

@app.post("/api/v1/convert-doc-to-xls")
async def convert_doc_xls(file: UploadFile = File(...)):
    import pandas as pd
    import docx
    def deleteFiles():
        # Si la carpeta pdf_files existe, elimina todos los archivos PDF y luego elimina la carpeta pdf_files
        if os.path.exists("doc_to_xls_files"):
            for doc_file_ in os.listdir("doc_to_xls_files"):
                os.remove(os.path.join("doc_to_xls_files", doc_file_))
            os.rmdir("doc_to_xls_files")
        # Si la carpeta csv_files existe, elimina todos los archivos CSV y luego elimina la carpeta csv_files
        if os.path.exists("doc_xls_files"):
            for xls_file_ in os.listdir("doc_xls_files"):
                a = open(os.path.join("doc_xls_files", xls_file_), "rb")
                a = a.close()

                os.remove(os.path.join("doc_xls_files", xls_file_))
            os.rmdir("doc_xls_files")
    deleteFiles()
    # Crea una carpeta para almacenar el PDF recibido
    os.mkdir("doc_to_xls_files")
    # Guardar el PDF recibido en el directorio pdf_files
    file_name = os.path.join("doc_to_xls_files", file.filename)
    # Abre el archivo PDF en modo binario y lo guarda en el directorio pdf_files
    with open(file_name, "wb") as buffer:
        # Lee el archivo PDF en trozos de 1024 bytes
        while chunk := await file.read(1024):
            # Escribe el trozo de bytes en el archivo PDF
            buffer.write(chunk)
    os.makedirs("doc_xls_files", exist_ok=True)


    # Cargar el archivo de Word
    doc = docx.Document(file_name)

    # Inicializar una lista vacía para almacenar los datos de la tabla
    table_data = []

    # Iterar por cada tabla en el documento
    for table in doc.tables:
        # Iterar por cada fila de la tabla
        for row in table.rows:
            # Inicializar una lista vacía para almacenar los datos de la fila
            row_data = []
            # Iterar por cada celda de la fila
            for cell in row.cells:
                # Agregar el contenido de la celda a la lista de datos de la fila
                row_data.append(cell.text)
            # Agregar los datos de la fila a la lista de datos de la tabla
            table_data.append(row_data)

    # Convertir los datos de la tabla en un DataFrame de Pandas
    df = pd.DataFrame(table_data)

    # Guardar el DataFrame en un archivo de Excel
    df.to_excel("doc_xls_files/converted.xls", index=False)

    zip_file_name = "xls_file.zip"
    with zipfile.ZipFile(zip_file_name, "w") as zip_file:
        # Agrega todos los archivos DOC a la carpeta doc_xls_files
        for xls_file_name in os.listdir("doc_xls_files"):
            csv_file = os.path.join("doc_xls_files", xls_file_name)
            zip_file.write(csv_file, xls_file_name)
    # Mueve el archivo zip a la carpeta csv_files
    os.rename("xls_file.zip", "doc_xls_files/xls_file.zip")
    data = open("doc_xls_files/xls_file.zip", "rb")
    file_contents = data.read()
    data.close()
    response = StreamingResponse(BytesIO(file_contents), media_type="application/vnd.ms-excel", headers={"Content-Disposition": "attachment; filename=xls_files.zip"})
    deleteFiles()
    return response

@app.post("/api/v1/convert-doc-to-docx")
async def convert_doc_docx(file: UploadFile = File(...)):
    import docx
    def deleteFiles():
        # Si la carpeta doc_files existe, elimina todos los archivos DOC y luego elimina la carpeta doc_files
        if os.path.exists("doc_files"):
            for doc_file_ in os.listdir("doc_files"):
                os.remove(os.path.join("doc_files", doc_file_))
            os.rmdir("doc_files")
        # Si la carpeta docx_files existe, elimina todos los archivos DOCX y luego elimina la carpeta docx_files
        if os.path.exists("docx_files"):
            for docx_file_ in os.listdir("docx_files"):
                a = open(os.path.join("docx_files", docx_file_), "rb")
                a = a.close()

                os.remove(os.path.join("docx_files", docx_file_))
            os.rmdir("docx_files")
    deleteFiles()
    # Crea una carpeta para almacenar el archivo DOC recibido
    os.mkdir("doc_files")
    # Guardar el archivo DOC recibido en el directorio doc_files
    file_name = os.path.join("doc_files", file.filename)
    # Abre el archivo DOC en modo binario y lo guarda en el directorio doc_files
    with open(file_name, "wb") as buffer:
        # Lee el archivo DOC en trozos de 1024 bytes
        while chunk := await file.read(1024):
            # Escribe el trozo de bytes en el archivo DOC
            buffer.write(chunk)
    os.makedirs("docx_files", exist_ok=True)

    # Cargar el archivo de Word en formato DOC
    doc = docx.Document(file_name)

    # Guardar el archivo de Word en formato DOCX
    docx_file_name = os.path.join("docx_files", os.path.splitext(file.filename)[0] + ".docx")
    doc.save(docx_file_name)

    # Crear un archivo zip con el archivo DOCX
    zip_file_name = "docx_file.zip"
    with zipfile.ZipFile(zip_file_name, "w") as zip_file:
        # Agrega el archivo DOCX a la carpeta docx_files
        zip_file.write(docx_file_name, os.path.basename(docx_file_name))
    # Mueve el archivo zip a la carpeta docx_files
    os.rename("docx_file.zip", os.path.join("docx_files", "docx_file.zip"))
    data = open(os.path.join("docx_files", "docx_file.zip"), "rb")
    file_contents = data.read()
    data.close()
    response = StreamingResponse(BytesIO(file_contents), media_type="application/docx", headers={"Content-Disposition": "attachment; filename=docx_files.zip"})
    deleteFiles()
    return response

@app.post("/api/v1/convert-doc-to-pdf")
async def convert_doc_pdf(file: UploadFile = File(...)):
    from docx2pdf import convert
    import docx
    def deleteFiles():
        # Si la carpeta docx_files existe, elimina todos los archivos DOCX y luego elimina la carpeta docx_files
        if os.path.exists("doc_to_pdf_files"):
            for docx_file_ in os.listdir("doc_to_pdf_files"):
                os.remove(os.path.join("doc_to_pdf_files", docx_file_))
            os.rmdir("doc_to_pdf_files")
        # Si la carpeta doc_files existe, elimina todos los archivos DOC y luego elimina la carpeta doc_files
        if os.path.exists("doc_pdf_files"):
            for doc_file_ in os.listdir("doc_pdf_files"):
                a = open(os.path.join("doc_pdf_files", doc_file_), "rb")
                a = a.close()

                os.remove(os.path.join("doc_pdf_files", doc_file_))
            os.rmdir("doc_pdf_files")
    deleteFiles()
    # Crea una carpeta para almacenar el archivo DOCX recibido
    os.mkdir("doc_to_pdf_files")
    # Guardar el archivo DOCX recibido en el directorio docx_files
    file_name = os.path.join("doc_to_pdf_files", file.filename)
    # Abre el archivo DOCX en modo binario y lo guarda en el directorio docx_files
    with open(file_name, "wb") as buffer:
        # Lee el archivo DOCX en trozos de 1024 bytes
        while chunk := await file.read(1024):
            # Escribe el trozo de bytes en el archivo DOCX
            buffer.write(chunk)
    os.makedirs("doc_pdf_files", exist_ok=True)
    
    # Cargar el archivo de Word en formato DOC
    doc = docx.Document(file_name)

    # Guardar el archivo de Word en formato DOCX
    docx_file_name = os.path.join("doc_to_pdf_files", os.path.splitext(file.filename)[0] + ".docx")
    doc.save(docx_file_name)

    convert(docx_file_name, "doc_pdf_files/converted.pdf")

    zip_file_name = "doc_file.zip"
    with zipfile.ZipFile(zip_file_name, "w") as zip_file:
        # Agrega todos los archivos DOC a la carpeta doc_pdf_files
        for doc_file_name in os.listdir("doc_pdf_files"):
            doc_file = os.path.join("doc_pdf_files", doc_file_name)
            zip_file.write(doc_file, doc_file_name)
    # Mueve el archivo zip a la carpeta doc_files
    os.rename("doc_file.zip", "doc_pdf_files/doc_file.zip")
    data = open("doc_pdf_files/doc_file.zip", "rb")
    file_contents = data.read()
    data.close()
    response = StreamingResponse(BytesIO(file_contents), media_type="application/pdf", headers={"Content-Disposition": "attachment; filename=doc_files.zip"})
    deleteFiles()
    return response



# Final conversor de DOC


# Inicio conversor de DOCX

@app.post("/api/v1/convert-docx-to-csv")
async def convert_doc_csv(file: UploadFile = File(...)):
    import pandas as pd
    import docx
    def deleteFiles():
        # Si la carpeta pdf_files existe, elimina todos los archivos PDF y luego elimina la carpeta pdf_files
        if os.path.exists("docx_to_csv_files"):
            for docx_file_ in os.listdir("docx_to_csv_files"):
                os.remove(os.path.join("docx_to_csv_files", docx_file_))
            os.rmdir("docx_to_csv_files")
        # Si la carpeta csv_files existe, elimina todos los archivos CSV y luego elimina la carpeta csv_files
        if os.path.exists("doc_csv_files"):
            for csv_file_ in os.listdir("doc_csv_files"):
                a = open(os.path.join("doc_csv_files", csv_file_), "rb")
                a = a.close()

                os.remove(os.path.join("doc_csv_files", csv_file_))
            os.rmdir("doc_csv_files")
    deleteFiles()
    # Crea una carpeta para almacenar el PDF recibido
    os.mkdir("doc_to_csv_files")
    # Guardar el PDF recibido en el directorio pdf_files
    file_name = os.path.join("doc_to_csv_files", file.filename)
    # Abre el archivo PDF en modo binario y lo guarda en el directorio pdf_files
    with open(file_name, "wb") as buffer:
        # Lee el archivo PDF en trozos de 1024 bytes
        while chunk := await file.read(1024):
            # Escribe el trozo de bytes en el archivo PDF
            buffer.write(chunk)
    os.makedirs("doc_csv_files", exist_ok=True)


    # Cargar el archivo de Word
    doc = docx.Document(file_name)

    # Inicializar una lista vacía para almacenar los datos de la tabla
    table_data = []

    # Iterar por cada tabla en el documento
    for table in doc.tables:
        # Iterar por cada fila de la tabla
        for row in table.rows:
            # Inicializar una lista vacía para almacenar los datos de la fila
            row_data = []
            # Iterar por cada celda de la fila
            for cell in row.cells:
                # Agregar el contenido de la celda a la lista de datos de la fila
                row_data.append(cell.text)
            # Agregar los datos de la fila a la lista de datos de la tabla
            table_data.append(row_data)

    # Convertir los datos de la tabla en un DataFrame de Pandas
    df = pd.DataFrame(table_data)

    # Guardar el DataFrame en un archivo de Excel
    df.to_csv("doc_csv_files/converted.csv", index=False)

    zip_file_name = "csv_file.zip"
    with zipfile.ZipFile(zip_file_name, "w") as zip_file:
        # Agrega todos los archivos DOC a la carpeta doc_csv_files
        for pdf_file_name in os.listdir("doc_csv_files"):
            csv_file = os.path.join("doc_csv_files", pdf_file_name)
            zip_file.write(csv_file, pdf_file_name)
    # Mueve el archivo zip a la carpeta csv_files
    os.rename("csv_file.zip", "doc_csv_files/csv_file.zip")
    data = open("doc_csv_files/csv_file.zip", "rb")
    file_contents = data.read()
    data.close()
    response = StreamingResponse(BytesIO(file_contents), media_type="application/csv", headers={"Content-Disposition": "attachment; filename=csv_files.zip"})
    deleteFiles()
    return response

@app.post("/api/v1/convert-docx-to-doc")
async def convert_docx_doc(file: UploadFile = File(...)):
    import docx
    def deleteFiles():
        # Si la carpeta docx_files existe, elimina todos los archivos DOCX y luego elimina la carpeta docx_files
        if os.path.exists("doc_to_pdf_files"):
            for docx_file_ in os.listdir("doc_to_pdf_files"):
                os.remove(os.path.join("doc_to_pdf_files", docx_file_))
            os.rmdir("doc_to_pdf_files")
        # Si la carpeta doc_files existe, elimina todos los archivos DOC y luego elimina la carpeta doc_files
        if os.path.exists("doc_pdf_files"):
            for doc_file_ in os.listdir("doc_pdf_files"):
                a = open(os.path.join("doc_pdf_files", doc_file_), "rb")
                a = a.close()

                os.remove(os.path.join("doc_pdf_files", doc_file_))
            os.rmdir("doc_pdf_files")
    deleteFiles()
    # Crea una carpeta para almacenar el archivo DOCX recibido
    os.mkdir("doc_to_pdf_files")
    # Guardar el archivo DOCX recibido en el directorio docx_files
    file_name = os.path.join("doc_to_pdf_files", file.filename)
    # Abre el archivo DOCX en modo binario y lo guarda en el directorio docx_files
    with open(file_name, "wb") as buffer:
        # Lee el archivo DOCX en trozos de 1024 bytes
        while chunk := await file.read(1024):
            # Escribe el trozo de bytes en el archivo DOCX
            buffer.write(chunk)
    os.makedirs("doc_pdf_files", exist_ok=True)

    # Cargar el archivo de Word
    doc = docx.Document(file_name)

    # Guardar el archivo DOC
    doc.save("doc_pdf_files/converted.doc")

    zip_file_name = "doc_file.zip"
    with zipfile.ZipFile(zip_file_name, "w") as zip_file:
        # Agrega todos los archivos DOC a la carpeta doc_pdf_files
        for doc_file_name in os.listdir("doc_pdf_files"):
            doc_file = os.path.join("doc_pdf_files", doc_file_name)
            zip_file.write(doc_file, doc_file_name)
    # Mueve el archivo zip a la carpeta doc_files
    os.rename("doc_file.zip", "doc_pdf_files/doc_file.zip")
    data = open("doc_pdf_files/doc_file.zip", "rb")
    file_contents = data.read()
    data.close()
    response = StreamingResponse(BytesIO(file_contents), media_type="application/msword", headers={"Content-Disposition": "attachment; filename=doc_files.zip"})
    deleteFiles()
    return response

@app.post("/api/v1/convert-docx-to-pdf")
async def convert_docx_pdf(file: UploadFile = File(...)):
    from docx2pdf import convert
    import docx
    def deleteFiles():
        # Si la carpeta docx_files existe, elimina todos los archivos DOCX y luego elimina la carpeta docx_files
        if os.path.exists("docx_to_pdf_files"):
            for docx_file_ in os.listdir("docx_to_pdf_files"):
                os.remove(os.path.join("docx_to_pdf_files", docx_file_))
            os.rmdir("docx_to_pdf_files")
        # Si la carpeta docx_files existe, elimina todos los archivos DOC y luego elimina la carpeta docx_files
        if os.path.exists("docx_pdf_files"):
            for docx_file_ in os.listdir("docx_pdf_files"):
                a = open(os.path.join("docx_pdf_files", docx_file_), "rb")
                a = a.close()

                os.remove(os.path.join("docx_pdf_files", docx_file_))
            os.rmdir("docx_pdf_files")
    deleteFiles()
    # Crea una carpeta para almacenar el archivo DOCX recibido
    os.mkdir("docx_to_pdf_files")
    # Guardar el archivo DOCX recibido en el directorio docx_files
    file_name = os.path.join("docx_to_pdf_files", file.filename)
    # Abre el archivo DOCX en modo binario y lo guarda en el directorio docx_files
    with open(file_name, "wb") as buffer:
        # Lee el archivo DOCX en trozos de 1024 bytes
        while chunk := await file.read(1024):
            # Escribe el trozo de bytes en el archivo DOCX
            buffer.write(chunk)
    os.makedirs("docx_pdf_files", exist_ok=True)
    
    # Cargar el archivo de Word en formato DOC
    doc = docx.Document(file_name)

    # Guardar el archivo de Word en formato DOCX
    docx_file_name = os.path.join("docx_to_pdf_files", os.path.splitext(file.filename)[0] + ".docx")
    doc.save(docx_file_name)

    convert(docx_file_name, "docx_pdf_files/converted.pdf")

    zip_file_name = "docx_file.zip"
    with zipfile.ZipFile(zip_file_name, "w") as zip_file:
        # Agrega todos los archivos DOC a la carpeta docx_pdf_files
        for docx_file_name in os.listdir("docx_pdf_files"):
            docx_file = os.path.join("docx_pdf_files", docx_file_name)
            zip_file.write(docx_file, docx_file_name)
    # Mueve el archivo zip a la carpeta docx_files
    os.rename("docx_file.zip", "docx_pdf_files/docx_file.zip")
    data = open("docx_pdf_files/docx_file.zip", "rb")
    file_contents = data.read()
    data.close()
    response = StreamingResponse(BytesIO(file_contents), media_type="application/pdf", headers={"Content-Disposition": "attachment; filename=docx_files.zip"})
    deleteFiles()
    return response

@app.post("/api/v1/convert-docx-to-xlsx")
async def convert_docx_xlsx(file: UploadFile = File(...)):
    import pandas as pd
    import docx
    def deleteFiles():
        # Si la carpeta pdf_files existe, elimina todos los archivos PDF y luego elimina la carpeta pdf_files
        if os.path.exists("docx_to_xlsx_files"):
            for docx_file_ in os.listdir("docx_to_xlsx_files"):
                os.remove(os.path.join("docx_to_xlsx_files", docx_file_))
            os.rmdir("docx_to_xlsx_files")
        # Si la carpeta csv_files existe, elimina todos los archivos CSV y luego elimina la carpeta csv_files
        if os.path.exists("docx_xlsx_files"):
            for xlsx_file_ in os.listdir("docx_xlsx_files"):
                a = open(os.path.join("docx_xlsx_files", xlsx_file_), "rb")
                a = a.close()

                os.remove(os.path.join("docx_xlsx_files", xlsx_file_))
            os.rmdir("docx_xlsx_files")
    deleteFiles()
    # Crea una carpeta para almacenar el PDF recibido
    os.mkdir("docx_to_xlsx_files")
    # Guardar el PDF recibido en el directorio pdf_files
    file_name = os.path.join("docx_to_xlsx_files", file.filename)
    # Abre el archivo PDF en modo binario y lo guarda en el directorio pdf_files
    with open(file_name, "wb") as buffer:
        # Lee el archivo PDF en trozos de 1024 bytes
        while chunk := await file.read(1024):
            # Escribe el trozo de bytes en el archivo PDF
            buffer.write(chunk)
    os.makedirs("docx_xlsx_files", exist_ok=True)


    # Cargar el archivo de Word
    doc = docx.Document(file_name)

    # Inicializar una lista vacía para almacenar los datos de la tabla
    table_data = []

    # Iterar por cada tabla en el documento
    for table in doc.tables:
        # Iterar por cada fila de la tabla
        for row in table.rows:
            # Inicializar una lista vacía para almacenar los datos de la fila
            row_data = []
            # Iterar por cada celda de la fila
            for cell in row.cells:
                # Agregar el contenido de la celda a la lista de datos de la fila
                row_data.append(cell.text)
            # Agregar los datos de la fila a la lista de datos de la tabla
            table_data.append(row_data)

    # Convertir los datos de la tabla en un DataFrame de Pandas
    df = pd.DataFrame(table_data)

    # Guardar el DataFrame en un archivo de Excel
    df.to_excel("docx_xlsx_files/converted.xlsx", index=False)

    zip_file_name = "xlsx_file.zip"
    with zipfile.ZipFile(zip_file_name, "w") as zip_file:
        # Agrega todos los archivos DOC a la carpeta doc_xlsx_files
        for pdf_file_name in os.listdir("docx_xlsx_files"):
            csv_file = os.path.join("docx_xlsx_files", pdf_file_name)
            zip_file.write(csv_file, pdf_file_name)
    # Mueve el archivo zip a la carpeta csv_files
    os.rename("xlsx_file.zip", "docx_xlsx_files/xlsx_file.zip")
    data = open("docx_xlsx_files/xlsx_file.zip", "rb")
    file_contents = data.read()
    data.close()
    response = StreamingResponse(BytesIO(file_contents), media_type="application/xlsx", headers={"Content-Disposition": "attachment; filename=csv_files.zip"})
    deleteFiles()
    return response

@app.post("/api/v1/convert-docx-to-xls")
async def convert_docx_xls(file: UploadFile = File(...)):
    import pandas as pd
    import docx
    def deleteFiles():
        # Si la carpeta pdf_files existe, elimina todos los archivos PDF y luego elimina la carpeta pdf_files
        if os.path.exists("doc_to_xls_files"):
            for docx_file_ in os.listdir("doc_to_xls_files"):
                os.remove(os.path.join("doc_to_xls_files", docx_file_))
            os.rmdir("doc_to_xls_files")
        # Si la carpeta csv_files existe, elimina todos los archivos CSV y luego elimina la carpeta csv_files
        if os.path.exists("doc_xls_files"):
            for xls_file_ in os.listdir("doc_xls_files"):
                a = open(os.path.join("doc_xls_files", xls_file_), "rb")
                a = a.close()

                os.remove(os.path.join("doc_xls_files", xls_file_))
            os.rmdir("doc_xls_files")
    deleteFiles()
    # Crea una carpeta para almacenar el PDF recibido
    os.mkdir("doc_to_xls_files")
    # Guardar el PDF recibido en el directorio pdf_files
    file_name = os.path.join("doc_to_xls_files", file.filename)
    # Abre el archivo PDF en modo binario y lo guarda en el directorio pdf_files
    with open(file_name, "wb") as buffer:
        # Lee el archivo PDF en trozos de 1024 bytes
        while chunk := await file.read(1024):
            # Escribe el trozo de bytes en el archivo PDF
            buffer.write(chunk)
    os.makedirs("doc_xls_files", exist_ok=True)


    # Cargar el archivo de Word
    doc = docx.Document(file_name)

    # Inicializar una lista vacía para almacenar los datos de la tabla
    table_data = []

    # Iterar por cada tabla en el documento
    for table in doc.tables:
        # Iterar por cada fila de la tabla
        for row in table.rows:
            # Inicializar una lista vacía para almacenar los datos de la fila
            row_data = []
            # Iterar por cada celda de la fila
            for cell in row.cells:
                # Agregar el contenido de la celda a la lista de datos de la fila
                row_data.append(cell.text)
            # Agregar los datos de la fila a la lista de datos de la tabla
            table_data.append(row_data)

    # Convertir los datos de la tabla en un DataFrame de Pandas
    df = pd.DataFrame(table_data)

    # Guardar el DataFrame en un archivo de Excel
    df.to_excel("doc_xls_files/converted.xls", index=False)

    zip_file_name = "xls_file.zip"
    with zipfile.ZipFile(zip_file_name, "w") as zip_file:
        # Agrega todos los archivos DOC a la carpeta doc_xls_files
        for xls_file_name in os.listdir("doc_xls_files"):
            csv_file = os.path.join("doc_xls_files", xls_file_name)
            zip_file.write(csv_file, xls_file_name)
    # Mueve el archivo zip a la carpeta csv_files
    os.rename("xls_file.zip", "doc_xls_files/xls_file.zip")
    data = open("doc_xls_files/xls_file.zip", "rb")
    file_contents = data.read()
    data.close()
    response = StreamingResponse(BytesIO(file_contents), media_type="application/xls", headers={"Content-Disposition": "attachment; filename=xls_files.zip"})
    deleteFiles()
    return response


@app.post("/api/v1/convert-zip-to-pdf")
async def convert_html_pdf(file: UploadFile = File(...)):
    import zipfile
    import tempfile
    import os
    import pdfkit
    def deleteFiles():
        # Si la carpeta pdf_files existe, elimina todos los archivos PDF y luego elimina la carpeta pdf_files
        if os.path.exists("html_to_pdf_files"):
            for docx_file_ in os.listdir("html_to_pdf_files"):
                os.remove(os.path.join("html_to_pdf_files", docx_file_))
            os.rmdir("html_to_pdf_files")
        # Si la carpeta csv_files existe, elimina todos los archivos CSV y luego elimina la carpeta csv_files
        if os.path.exists("html_pdf_files"):
            for html_file_ in os.listdir("html_pdf_files"):
                a = open(os.path.join("html_pdf_files", html_file_), "rb")
                a = a.close()

                os.remove(os.path.join("html_pdf_files", html_file_))
            os.rmdir("html_pdf_files")
    deleteFiles()
    # Crea una carpeta para almacenar el PDF recibido
    os.mkdir("html_to_pdf_files")
    # Guardar el PDF recibido en el directorio pdf_files
    file_name = os.path.join("html_to_pdf_files", file.filename)
    # Abre el archivo PDF en modo binario y lo guarda en el directorio pdf_files
    with open(file_name, "wb") as buffer:
        # Lee el archivo PDF en trozos de 1024 bytes
        while chunk := await file.read(1024):
            # Escribe el trozo de bytes en el archivo PDF
            buffer.write(chunk)
    os.makedirs("html_pdf_files", exist_ok=True)
    

    # Extract the zipped website to a temporary directory
    with zipfile.ZipFile(file_name, 'r') as zip_ref:
        zip_ref.extractall("html_to_pdf_files")

    # Convert the HTML file to a PDF using pdfkit
    for pdf_file_name in os.listdir("html_to_pdf_files"):
        if pdf_file_name.endswith(".html"):
            html_file = os.path.join("html_to_pdf_files", pdf_file_name)
    pdfkit.from_file(html_file, 'html_pdf_files/converted.pdf')

    zip_file_name = "pdf_file.zip"
    with zipfile.ZipFile(zip_file_name, "w") as zip_file:
        # Agrega todos los archivos DOC a la carpeta html_pdf_files
        for pdf_file_name in os.listdir("html_pdf_files"):
            html_file = os.path.join("html_pdf_files", pdf_file_name)
            zip_file.write(html_file, pdf_file_name)
    # Mueve el archivo zip a la carpeta csv_files
    os.rename("pdf_file.zip", "html_pdf_files/pdf_file.zip")
    data = open("html_pdf_files/pdf_file.zip", "rb")
    file_contents = data.read()
    data.close()
    response = StreamingResponse(BytesIO(file_contents), media_type="application/pdf", headers={"Content-Disposition": "attachment; filename=html_files.zip"})
    deleteFiles()
    return response